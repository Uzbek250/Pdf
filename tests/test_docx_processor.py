"""
Test 2: DOCX run-darajasidagi tarjima taqsimoti — formatlash (bold/italic)
run obyektlarida saqlanib qolishini tekshiradi.
"""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from services.docx_processor import (
    _collect_translatable_paragraphs,
    _distribute_translation_to_runs,
    _get_all_runs_including_hyperlinks,
    _group_runs_by_format,
    _run_contains_drawing,
    _run_contains_footnote_reference,
)


def _make_multi_run_paragraph():
    """Bitta paragrafda 2 ta run (bold va oddiy) yaratadi."""
    document = Document()
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run("Hello ")
    run1.bold = True
    run2 = paragraph.add_run("world")
    run2.italic = True
    return document, paragraph, run1, run2


def test_distribute_translation_preserves_run_formatting_flags() -> None:
    """Tarjimadan keyin run.bold / run.italic o'zgarishsiz qolishi kerak."""
    document, paragraph, run1, run2 = _make_multi_run_paragraph()
    assert paragraph.text == "Hello world"

    _distribute_translation_to_runs(paragraph, "Salom dunyo")

    # Format bayroqlari saqlanishi kerak (run obyektlari o'zgarmagan)
    assert paragraph.runs[0].bold is True
    assert paragraph.runs[1].italic is True

    # Matn to'liq va to'g'ri taqsimlangan (ikkala run yig'indisi)
    combined = "".join(run.text for run in paragraph.runs)
    assert combined == "Salom dunyo"


def test_distribute_translation_single_run_replaces_text_directly() -> None:
    """Faqat bitta run bo'lganda, matn to'g'ridan-to'g'ri almashtiriladi."""
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Original matn")
    run.font.size = None  # placeholder, format tekshirilmaydi bu yerda

    _distribute_translation_to_runs(paragraph, "Tarjima qilingan matn")

    assert len(paragraph.runs) == 1
    assert paragraph.runs[0].text == "Tarjima qilingan matn"


def test_distribute_translation_total_length_matches_translated_text() -> None:
    """Ko'p run holatida ham umumiy uzunlik tarjima matni uzunligiga teng bo'lishi kerak
    (yaxlitlash xatolari to'g'rilanganini tekshiradi)."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A")  # juda qisqa run
    paragraph.add_run("BB")
    paragraph.add_run("CCC")

    translated = "Bu ancha uzunroq tarjima qilingan matn bo'lishi mumkin"
    _distribute_translation_to_runs(paragraph, translated)

    combined = "".join(run.text for run in paragraph.runs)
    assert combined == translated
    assert len(combined) == len(translated)


def test_collect_translatable_paragraphs_skips_empty_paragraphs() -> None:
    """Bo'sh paragraflar tarjima ro'yxatiga kiritilmasligi kerak."""
    document = Document()
    document.add_paragraph("Birinchi paragraf")
    document.add_paragraph("")  # bo'sh
    document.add_paragraph("   ")  # faqat bo'shliq
    document.add_paragraph("Ikkinchi paragraf")

    refs = _collect_translatable_paragraphs(document)
    texts = [ref.original_text for ref in refs]

    assert texts == ["Birinchi paragraf", "Ikkinchi paragraf"]


def test_distribute_translation_preserves_word_boundaries() -> None:
    """Tarjima matni run'lar orasiga bo'linganda so'zlar orasidagi
    bo'shliqlar yo'qolmasligi kerak (masalan "Matnning yangi" ->
    "Matnningyangi" kabi xatoga yo'l qo'yilmaydi)."""
    document = Document()
    paragraph = document.add_paragraph()
    # Ko'p sonli qisqa run'lar — bu Word imlo tekshiruvi natijasida
    # ko'pincha yuzaga keladigan holatni simulyatsiya qiladi.
    paragraph.add_run("Matn")
    paragraph.add_run("ning")
    paragraph.add_run(" yangi")
    paragraph.add_run(" leksik")

    translated = "Text's new lexical"
    _distribute_translation_to_runs(paragraph, translated)

    combined = "".join(run.text for run in paragraph.runs)
    # Umumiy matn to'g'ri va HECH QANDAY so'z chegarasi buzilmagan bo'lishi
    # kerak — ya'ni combined ichida so'zlar orasida bo'shliq saqlanadi.
    assert combined == translated
    assert "  " not in combined  # ikki marta bo'shliq ketma-ket kelmasligi kerak
    # So'zlar orasida bo'shliq yo'qolib, ular qo'shilib ketmagan bo'lishi kerak
    assert "Text'snew" not in combined
    assert "newlexical" not in combined


def test_distribute_translation_skips_runs_containing_drawing() -> None:
    """Rasm (drawing) saqlovchi run'lar tarjima paytida tegilmasligi va
    ularning ichidagi rasm elementi yo'qolmasligi kerak."""
    document = Document()
    paragraph = document.add_paragraph()
    text_run = paragraph.add_run("Original matn")

    # Rasmni simulyatsiya qilish uchun ikkinchi run yaratib, uning XML
    # elementiga qo'lda soxta <w:drawing> elementini qo'shamiz — bu
    # haqiqiy rasm run'ining tuzilishini taqlid qiladi.
    image_run = paragraph.add_run("")
    drawing_xml = (
        '<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "</w:drawing>"
    )
    from lxml import etree

    drawing_element = etree.fromstring(drawing_xml)
    image_run._element.append(drawing_element)

    assert _run_contains_drawing(image_run) is True
    assert _run_contains_drawing(text_run) is False

    _distribute_translation_to_runs(paragraph, "Tarjima qilingan matn")

    # Matn run'i tarjima qilingan bo'lishi kerak
    assert text_run.text == "Tarjima qilingan matn"
    # Rasm run'ining matni o'zgarishsiz (bo'sh) qolishi va rasm elementi
    # hali ham mavjud bo'lishi kerak.
    assert image_run.text == ""
    assert _run_contains_drawing(image_run) is True


def test_distribute_translation_skips_footnote_reference_runs() -> None:
    """Footnote/endnote referens belgisi saqlovchi run'lar tarjima
    paytida tegilmasligi kerak — aks holda footnote raqami yo'qolib
    qolishi mumkin."""
    document = Document()
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run("Matnning boshi ")

    # Footnote referensini simulyatsiya qilamiz.
    footnote_run = paragraph.add_run("")
    footnote_xml = (
        '<w:footnoteReference xmlns:w='
        '"http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:id="1"/>'
    )
    from lxml import etree

    footnote_element = etree.fromstring(footnote_xml)
    footnote_run._element.append(footnote_element)

    run2 = paragraph.add_run(" va matnning davomi")

    assert _run_contains_footnote_reference(footnote_run) is True
    assert _run_contains_footnote_reference(run1) is False

    _distribute_translation_to_runs(
        paragraph, "Tarjima qilingan matnning boshi va davomi"
    )

    # Footnote run'i tegilmagan (bo'sh) qolishi kerak.
    assert footnote_run.text == ""
    assert _run_contains_footnote_reference(footnote_run) is True

    # Matn run'lari (footnote'dan tashqari) tarjima matnini olishi kerak,
    # va ular birlashganda to'liq tarjima matni chiqishi kerak.
    combined = run1.text + footnote_run.text + run2.text
    assert combined == "Tarjima qilingan matnning boshi va davomi"


def test_group_runs_by_format_merges_same_format_consecutive_runs() -> None:
    """Bir xil formatga ega, ketma-ket kelgan run'lar bitta guruhga
    birlashishi kerak — bu ba'zi hujjatlarda uchraydigan "sun'iy" run
    bo'linishini (bitta so'zning bir necha run'ga bo'linib ketishi)
    tuzatish uchun zarur."""
    document = Document()
    paragraph = document.add_paragraph()
    # Xuddi shu formatdagi (hammasi oddiy, formatlashsiz) 3 ta run —
    # bular bitta guruhga birlashishi kerak.
    r1 = paragraph.add_run("орга")
    r2 = paragraph.add_run("низма")
    r3 = paragraph.add_run(" татко")
    # Formatlash farqi bo'lgan run — alohida guruh bo'lishi kerak.
    r4 = paragraph.add_run("vy")
    r4.bold = True

    groups = _group_runs_by_format(paragraph.runs)

    assert len(groups) == 2  # birinchi 3 tasi bitta guruh, oxirgisi alohida
    assert [r.text for r in groups[0]] == ["орга", "низма", " татко"]
    assert [r.text for r in groups[1]] == ["vy"]


def _add_hyperlink_run(paragraph: Paragraph, text: str) -> None:
    """Test uchun paragrafga XML orqali hyperlink elementi qo'shadi.

    python-docx'da hyperlink qo'shish uchun yuqori darajali ochiq API
    yo'q, shuning uchun past darajali `makeelement` orqali quramiz — bu
    `docx.oxml` registry'siga to'g'ri bog'lanishini ta'minlaydi (aks
    holda `paragraph.text` xususiyati ichki xato berishi mumkin).
    """
    hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {})
    run_el = hyperlink.makeelement(qn("w:r"), {})
    t_el = run_el.makeelement(qn("w:t"), {})
    t_el.text = text
    run_el.append(t_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def test_get_all_runs_including_hyperlinks_finds_hyperlink_text() -> None:
    """Paragraf ichidagi <w:hyperlink> elementiga joylashgan matn ham
    run ro'yxatiga kiritilishi kerak.

    MUHIM REAL MUAMMO: `paragraph.runs` (python-docx standart xususiyati)
    faqat paragrafning BEVOSITA farzand `<w:r>` elementlarini qaytaradi.
    Agar paragrafda gipermatn (hyperlink, masalan "batafsil ma'lumot"
    kabi havola) bo'lsa, uning matni `<w:hyperlink><w:r>...</w:r></w:hyperlink>`
    ichida joylashadi va oddiy `paragraph.runs` uni KO'RMAYDI — garchi
    `paragraph.text` uni avtomatik hisoblasa ham. Bu nomuvofiqlik
    (paragraph.text hyperlink matnini o'z ichiga oladi, lekin
    paragraph.runs olmaydi) haqiqiy hujjatda gipermatn so'zining tarjima
    natijasidan tushib qolishiga yoki gap oxiriga yopishib qolishiga olib
    kelgan edi."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Ma'lumot uchun ")
    _add_hyperlink_run(paragraph, "arxivga")
    paragraph.add_run(" qarang.")

    # `paragraph.text` hyperlink matnini avtomatik qo'shib hisoblaydi.
    assert paragraph.text == "Ma'lumot uchun arxivga qarang."

    # Ammo standart `paragraph.runs` uni KO'RMAYDI (bu — muammoning o'zi).
    assert len(paragraph.runs) == 2  # faqat 2 ta oddiy run, hyperlink emas

    # Bizning yordamchi funksiyamiz esa hyperlink ichidagi matnni ham
    # to'g'ri topishi kerak.
    all_runs = _get_all_runs_including_hyperlinks(paragraph)
    assert len(all_runs) == 3
    assert [r.text for r in all_runs] == ["Ma'lumot uchun ", "arxivga", " qarang."]


def test_distribute_translation_handles_hyperlink_text_without_losing_content() -> None:
    """Hyperlink matni bo'lgan paragrafda tarjima qilinganda, umumiy
    matn mazmuni to'liq saqlanishi kerak (hyperlink matni yo'qolmasligi
    yoki gap oxiriga yopishib qolmasligi kerak)."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("See the ")
    _add_hyperlink_run(paragraph, "archive")
    paragraph.add_run(" for details.")

    translated = "Batafsil ma'lumot uchun arxivga qarang."
    _distribute_translation_to_runs(paragraph, translated)

    all_runs_after = _get_all_runs_including_hyperlinks(paragraph)
    combined = "".join(r.text for r in all_runs_after)

    # Eng muhim tekshiruv: hyperlink borligiga qaramay, tarjima matni
    # TO'LIQ va TO'G'RI bo'lishi kerak — hech qanday so'z yo'qolmasligi
    # yoki gap oxiriga yopishib qolmasligi kerak.
    assert combined == translated
