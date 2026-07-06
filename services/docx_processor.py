"""
DOCX hujjatlarini RUN darajasida tarjima qilish.

MUHIM ARXITEKTURAVIY QOIDA:
    python-docx'da `paragraph.text` faqat o'qish uchun qulay, lekin unga
    yozish (`paragraph.text = "..."`) BARCHA run'larni (bold, italic,
    rang, shrift kabi formatlashni) yo'qotib, bitta yagona run bilan
    almashtiradi. Natijada original hujjat formati butunlay buziladi.

    Shu sabab biz doim `run.text` darajasida ishlaymiz: har bir paragraf
    o'z run'lariga bo'linadi, faqat matn qismi tarjima qilinadi, format
    xususiyatlari (run.bold, run.italic, run.font, run.color va h.k.)
    umuman tegilmaydi.

    Muammoli holat: bitta jumla ko'pincha bir necha run'ga bo'linib
    ketadi (masalan, Word avtomatik imlo tekshiruvi tufayli). Agar har
    bir run'ni alohida tarjima qilsak, ma'no yo'qoladi (masalan "Hello"
    va " world" alohida-alohida tarjima qilinsa noto'g'ri natija chiqadi).
    Shu sababli biz avval paragraf run'larini "matn segmentlariga"
    yig'amiz, BUTUN paragraf matnini tarjima qilamiz, so'ngra tarjima
    qilingan matnni run'lar orasiga so'z chegaralarini hurmat qilgan
    holda qayta taqsimlaymiz — shu orqali formatlash imkon qadar
    saqlanadi va so'zlar o'rtasidagi bo'shliqlar yo'qolmaydi.

MUHIM: RASMLI RUN'LAR
    Bitta `Run` obyekti nafaqat oddiy matn, balki inline rasm
    (`<w:drawing>` yoki eski uslub `<w:pict>` XML elementi) ham bo'lishi
    mumkin. Agar bunday run'ning `.text` xususiyatiga yozsak,
    python-docx uning butun XML mazmunini (jumladan rasm elementini ham)
    o'chirib, faqat oddiy matn bilan almashtiradi — bu rasmni butunlay
    yo'q qilib yuboradi. Shu sabab biz har bir run'ni tarjima matni bilan
    to'ldirishdan oldin, u rasm saqlovchi run emasligini tekshiramiz va
    rasmli run'larni umuman tegmasdan qoldiramiz.
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from services.translator import TranslatorService

logger = logging.getLogger(__name__)

# Rasm yoki boshqa grafik obyektlarni ifodalovchi XML teglar (namespace
# prefiksidan qat'iy nazar, shuning uchun `endswith` orqali tekshiramiz).
_DRAWING_TAG_SUFFIXES = ("}drawing", "}pict", "}object")

# Footnote/endnote referens belgilarini (masalan matndagi kichik "¹" raqami)
# ifodalovchi XML teglar. Bular oddiy matn emas — Word ularni alohida
# hisoblaydi va avtomatik raqamlaydi. Agar bunday run'ning matnini
# o'zgartirsak yoki uni "oddiy matn run'i" sifatida hisoblab tarjima
# taqsimotiga aralashtirsak, footnote belgisi yo'qolishi yoki noto'g'ri
# joyga tushib qolishi mumkin.
_FOOTNOTE_REFERENCE_TAG_SUFFIXES = (
    "}footnoteReference",
    "}endnoteReference",
    "}footnoteRef",
    "}endnoteRef",
)


@dataclass(slots=True)
class ParagraphRef:
    """Bitta paragrafga va uning joriy matniga havola (tarjima navbati uchun)."""

    paragraph: Paragraph
    original_text: str
    run_lengths: list[int] = field(default_factory=list)


def _get_all_runs_including_hyperlinks(paragraph: Paragraph) -> list[Run]:
    """Paragrafdagi BARCHA matn run'larini qaytaradi — jumladan
    `<w:hyperlink>` ichiga joylashgan run'larni ham.

    MUHIM: `paragraph.runs` (python-docx standart xususiyati) faqat
    paragrafning BEVOSITA farzand elementlarini (`<w:r>`) qaytaradi.
    Hyperlink matni esa Word'da `<w:hyperlink><w:r>...</w:r></w:hyperlink>`
    ko'rinishida saqlanadi — bu holda `<w:r>` hyperlink elementining
    farzandi, paragrafning emas, shuning uchun `paragraph.runs` uni
    ko'rmaydi.

    Ammo `paragraph.text` (o'qish uchun property) hyperlink matnini
    AVTOMATIK qo'shib hisoblaydi. Natijada, agar biz tarjima uchun
    `paragraph.text`dan foydalanib, taqsimlash uchun `paragraph.runs`dan
    foydalansak, ikkalasi orasida nomuvofiqlik yuzaga keladi — hyperlink
    matni "yo'qolgandek" yoki noto'g'ri joyga qo'shilib qolgandek
    ko'rinadi (masalan gap oxiriga yopishib qoladi).

    Bu funksiya paragraf XML daraxtini chuqurroq aylanib chiqib, hyperlink
    ichidagi run'larni ham asosiy ro'yxatga, ular hujjatda qanday tartibda
    joylashgan bo'lsa shu tartibda qo'shadi.
    """
    runs: list[Run] = []
    for child in paragraph._element:
        tag = child.tag
        if tag == qn("w:r"):
            runs.append(Run(child, paragraph))
        elif tag == qn("w:hyperlink"):
            for grandchild in child:
                if grandchild.tag == qn("w:r"):
                    runs.append(Run(grandchild, paragraph))
    return runs


def _run_contains_drawing(run: Run) -> bool:
    """Berilgan run ichida rasm/chizma (drawing) elementi bormi tekshiradi.

    Bunday run'ning `.text` xususiyatiga yozish uning ichidagi rasmni
    yo'q qilib yuboradi, shuning uchun bunday run'larni tarjima matnini
    taqsimlashda umuman chetlab o'tamiz.
    """
    for child in run._element.iter():
        if any(child.tag.endswith(suffix) for suffix in _DRAWING_TAG_SUFFIXES):
            return True
    return False


def _run_contains_footnote_reference(run: Run) -> bool:
    """Berilgan run ichida footnote/endnote referens belgisi bormi tekshiradi.

    Bunday run'lar odatda bo'sh yoki juda qisqa matnga ega (Word ularni
    o'zi avtomatik raqamlaydi), lekin ular MUHIM strukturaviy elementlar —
    ularning matnini tarjima taqsimotiga aralashtirish footnote belgisini
    yo'qotib yuborishi yoki noto'g'ri joyga surib qo'yishi mumkin. Shu
    sabab bunday run'lar ham, xuddi rasmli run'lar kabi, tarjima matni
    taqsimlanishidan chetlab o'tiladi.
    """
    for child in run._element.iter():
        if any(
            child.tag.endswith(suffix)
            for suffix in _FOOTNOTE_REFERENCE_TAG_SUFFIXES
        ):
            return True
    return False


def _iter_block_paragraphs(document: Document):
    """Hujjatdagi barcha paragraflarni (asosiy tana + jadval kataklari) yig'adi.

    Header/footer alohida qayta ishlanadi (docx_processor.translate_docx ichida).
    """
    yield from document.paragraphs

    for table in document.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _collect_translatable_paragraphs(document: Document) -> list[ParagraphRef]:
    """Tarjima qilinishi kerak bo'lgan paragraflarni yig'adi (bo'shlarni o'tkazib yuboradi)."""
    refs: list[ParagraphRef] = []
    for paragraph in _iter_block_paragraphs(document):
        text = paragraph.text
        if not text or not text.strip():
            continue
        all_runs = _get_all_runs_including_hyperlinks(paragraph)
        run_lengths = [len(run.text) for run in all_runs]
        # Agar run'lar umuman bo'lmasa (kamdan-kam, lekin nazariy jihatdan
        # mumkin), paragraph.text'ga tayanamiz va keyin bitta "virtual run"
        # sifatida ko'ramiz.
        if not run_lengths:
            run_lengths = [len(text)]
        refs.append(ParagraphRef(paragraph=paragraph, original_text=text, run_lengths=run_lengths))
    return refs


def _run_format_signature(run: Run) -> tuple:
    """Run formatlashini solishtirish uchun taqqoslanadigan "imzo" tuzadi.

    Ikkita run bir xil imzoga ega bo'lsa, ular vizual jihatdan bir xil
    ko'rinadi (bold/italic/underline/shrift/rang/o'lcham bir xil), shuning
    uchun ularni bitta "samarali run" sifatida birlashtirish xavfsiz —
    formatlash yo'qolmaydi.
    """
    font = run.font
    color = None
    if font.color is not None and font.color.type is not None:
        color = str(font.color.rgb) if font.color.rgb else str(font.color.type)
    return (
        run.bold,
        run.italic,
        run.underline,
        font.name,
        font.size,
        color,
        font.strike,
        font.subscript,
        font.superscript,
    )


def _group_runs_by_format(runs: list[Run]) -> list[list[Run]]:
    """Ketma-ket kelgan, bir xil formatga ega run'larni guruhlarga ajratadi.

    MUHIM: bu ba'zi hujjatlarda (masalan skanerlangan yoki eski Word
    fayllardan konvertatsiya qilingan hujjatlarda) uchraydigan holatni
    tuzatish uchun kerak — bunday hujjatlarda bitta so'zning o'zi bir
    necha run'ga bo'linib ketgan bo'lishi mumkin (masalan "организма"
    so'zi "орга" + "низма" kabi ikkita alohida run sifatida saqlangan,
    garchi ular orasida hech qanday format farqi bo'lmasa ham).

    Bunday "sun'iy" run bo'linishlarini hisobga olmasdan tarjima matnini
    run soniga qarab taqsimlasak, tarjima matni ham xuddi shu tasodifiy
    joylardan kesiladi va so'zlar buzilib qoladi (masalan bo'shliqlar
    noto'g'ri joyga tushib qoladi).

    Yechim: agar ketma-ket kelgan run'lar bir xil formatga ega bo'lsa,
    ularni bitta guruh deb hisoblaymiz — tarjima matni shu guruhlar
    soniga qarab taqsimlanadi, guruh ichidagi run'lardan birortasiga
    (odatda birinchisiga) yoziladi, qolganlari bo'sh qilinadi. Bu orqali
    tarjima matni tasodifiy so'z-ichi chegaralarga emas, balki haqiqiy
    formatlash chegaralariga mos taqsimlanadi.
    """
    if not runs:
        return []

    groups: list[list[Run]] = [[runs[0]]]
    last_signature = _run_format_signature(runs[0])

    for run in runs[1:]:
        signature = _run_format_signature(run)
        if signature == last_signature:
            groups[-1].append(run)
        else:
            groups.append([run])
            last_signature = signature

    return groups


def _split_translated_text_by_word_boundaries(
    translated_text: str, num_targets: int
) -> list[str]:
    """Tarjima matnini so'z chegaralarini hurmat qilgan holda `num_targets`
    ta segmentga bo'ladi.

    Oddiy nisbiy-uzunlik bo'yicha kesish so'z o'rtasidan kesib, bo'shliqni
    yo'qotib yuborishi mumkin (masalan "Matnning yangi" -> "Matnningyangi").
    Bu funksiya avval matnni so'zlarga (bo'shliqni saqlagan holda) ajratadi,
    so'ng har bir segmentga nisbatan teng miqdorda so'z guruhini beradi —
    shunda bo'shliqlar hech qachon segment ichida yo'qolmaydi.
    """
    if num_targets <= 1:
        return [translated_text]

    # Bo'shliqlarni ham alohida "token" sifatida saqlab, matnni bo'lakларга
    # ajratamiz — shunda qayta birlashtirilganda hech narsa yo'qolmaydi.
    tokens = re.findall(r"\S+|\s+", translated_text)
    if not tokens:
        return [""] * num_targets

    total_len = len(translated_text)
    target_len_per_segment = total_len / num_targets

    segments: list[str] = []
    current_segment: list[str] = []
    current_len = 0
    remaining_targets = num_targets

    for idx, token in enumerate(tokens):
        current_segment.append(token)
        current_len += len(token)

        remaining_tokens = tokens[idx + 1 :]
        is_last_target = remaining_targets == 1

        if is_last_target:
            # Oxirgi segment — qolgan barcha tokenlarni oladi.
            continue

        if current_len >= target_len_per_segment and not token.isspace():
            # So'z chegarasida (bo'shliq emas) to'xtaymiz, shunda so'z
            # o'rtasidan bo'linmaydi.
            segments.append("".join(current_segment))
            current_segment = []
            current_len = 0
            remaining_targets -= 1
            if not remaining_tokens:
                break

    # Qolgan tokenlarni oxirgi segmentga qo'shamiz.
    if current_segment:
        segments.append("".join(current_segment))

    # Yetarli segment hosil bo'lmagan bo'lsa (juda qisqa matn holatida),
    # bo'sh segmentlar bilan to'ldiramiz.
    while len(segments) < num_targets:
        segments.append("")

    # Agar biror sababdan ortiqcha segment hosil bo'lsa (kamdan-kam holat),
    # ortiqchalarini oxirgisiga birlashtiramiz.
    if len(segments) > num_targets:
        extra = "".join(segments[num_targets - 1 :])
        segments = segments[: num_targets - 1] + [extra]

    return segments


def _distribute_translation_to_runs(paragraph: Paragraph, translated_text: str) -> None:
    """Tarjima qilingan matnni paragraf run'lariga taqsimlaydi.

    Ikki bosqichli himoya + guruhlash strategiyasi ishlatiladi:

    1. RASM va FOOTNOTE REFERENSLARI TEGILMAYDI: bunday run'lar tarjima
       taqsimotidan butunlay chetlab o'tiladi — ularning matni yozilmaydi,
       run obyektining o'zi ham o'chirilmaydi. Bu ularning ichidagi maxsus
       XML elementini (rasm yoki footnote raqami) yo'qotib qo'ymaslik
       uchun zarur.

    2. FORMATGA QARAB GURUHLASH: qolgan (oddiy matnli) run'lar orasida,
       ketma-ket kelgan va bir xil formatga (bold/italic/shrift/rang)
       ega bo'lganlari BITTA guruh sifatida ko'riladi. Bu ba'zi
       hujjatlarda (masalan eski yoki avtomatik konvertatsiya qilingan
       fayllarda) uchraydigan "sun'iy" run bo'linishini (bitta so'zning
       o'zi bir necha run'ga bo'linib ketishi) tuzatadi — aks holda
       tarjima matni ham xuddi shu tasodifiy joylardan kesilib, so'zlar
       yoki bo'shliqlar buzilib qolar edi.

    Tarjima matni shu guruhlar soniga qarab taqsimlanadi (so'z
    chegaralarini hurmat qilgan holda); har bir guruh ichida faqat
    BIRINCHI run matnni oladi, qolganlari bo'sh qilinadi — bu format
    xususiyatlarini (guruhning o'zi bir xil formatli bo'lgani uchun)
    yo'qotmaydi.
    """
    all_runs: list[Run] = _get_all_runs_including_hyperlinks(paragraph)
    if not all_runs:
        return

    # 1-bosqich: rasm yoki footnote/endnote referensi saqlovchi run'larni
    # ajratib olamiz — ularga umuman tegilmaydi.
    text_runs = [
        run
        for run in all_runs
        if not _run_contains_drawing(run) and not _run_contains_footnote_reference(run)
    ]

    if not text_runs:
        # Paragrafda faqat rasm/footnote bor, oddiy matn run'i yo'q.
        return

    if len(text_runs) == 1:
        text_runs[0].text = translated_text
        return

    # 2-bosqich: bir xil formatli ketma-ket run'larni guruhlaymiz —
    # shunda tarjima matni "sun'iy" run chegaralariga emas, balki
    # haqiqiy formatlash chegaralariga mos taqsimlanadi.
    groups = _group_runs_by_format(text_runs)

    if len(groups) == 1:
        # Barcha run'lar bir xil formatda — birinchisiga yozamiz,
        # qolganlarini bo'shatamiz.
        first_run, *rest_runs = groups[0]
        first_run.text = translated_text
        for run in rest_runs:
            run.text = ""
        return

    segments = _split_translated_text_by_word_boundaries(translated_text, len(groups))
    for group, segment in zip(groups, segments):
        first_run, *rest_runs = group
        first_run.text = segment
        for run in rest_runs:
            run.text = ""


def _translate_header_footer(
    document: Document, translations_map: dict[str, str]
) -> None:
    """Header/footer paragraflarini translations_map orqali tarjima qiladi.

    (translations_map allaqachon tarjima qilingan {original_text: translated_text}
    lug'ati — bu funksiya faqat qo'llaydi, yangi API so'rovi qilmaydi.)
    """
    for section in document.sections:
        for container in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if container is None:
                continue
            for paragraph in container.paragraphs:
                text = paragraph.text
                if text and text.strip() and text in translations_map:
                    _distribute_translation_to_runs(paragraph, translations_map[text])


async def translate_docx(
    input_path: Path,
    output_path: Path,
    target_language: str,
    translator: TranslatorService,
    source_language: str | None = None,
) -> Path:
    """DOCX faylini run-darajasida tarjima qilib, yangi faylga saqlaydi.

    Args:
        input_path: manba .docx fayli.
        output_path: tarjima qilingan .docx fayl saqlanadigan joy.
        target_language: maqsad til kodi.
        translator: TranslatorService instansi (provider + cache).
        source_language: agar oldindan aniqlangan bo'lsa, manba til kodi.

    Returns:
        output_path (qulaylik uchun).
    """
    document = Document(str(input_path))

    body_refs = _collect_translatable_paragraphs(document)
    unique_texts = list({ref.original_text for ref in body_refs})

    logger.info(
        "DOCX tarjima boshlandi: %s ta noyob paragraf (%s ta jami)",
        len(unique_texts),
        len(body_refs),
    )

    translated_list = await translator.translate_paragraphs(
        paragraphs=unique_texts,
        target_language=target_language,
        source_language=source_language,
    )
    translations_map: dict[str, str] = dict(zip(unique_texts, translated_list))

    for ref in body_refs:
        translated = translations_map.get(ref.original_text)
        if translated is None:
            continue
        _distribute_translation_to_runs(ref.paragraph, translated)

    # Header/footer'larni ham xuddi shu tarjima lug'atidan foydalanib qayta ishlaymiz.
    # Ularni alohida to'plab, kesh-miss bo'lganlarni qo'shimcha tarjima qilamiz.
    header_footer_texts: set[str] = set()
    for section in document.sections:
        for container in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if container is None:
                continue
            for paragraph in container.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    header_footer_texts.add(paragraph.text)

    missing_hf_texts = [t for t in header_footer_texts if t not in translations_map]
    if missing_hf_texts:
        hf_translated = await translator.translate_paragraphs(
            paragraphs=missing_hf_texts,
            target_language=target_language,
            source_language=source_language,
        )
        translations_map.update(dict(zip(missing_hf_texts, hf_translated)))

    _translate_header_footer(document, translations_map)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    logger.info("DOCX tarjima yakunlandi: %s", output_path)
    return output_path
